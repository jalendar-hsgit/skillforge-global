from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from typing import List, Optional
from datetime import datetime
from pydantic import BaseModel

from app.core.db import get_db
from app.core.security import get_current_user
from app.models.user import User
from app.modelsx.resume import Resume
from app.services.email_service import email_service
from app.core.config import settings
from app.services.llm_provider import get_provider
from app.modelsx.resume import ResumeTemplate
from sqlalchemy import or_

import os
import tempfile
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document

router = APIRouter(prefix="/resumes", tags=["resumes"])


@router.get("/{resume_id}/export_debug")
def export_resume_debug(resume_id: int, db: Session = Depends(get_db)):
    """Debug-only endpoint: return a small JSON summary of a resume for quick verification.
    Only available when `settings.DEBUG` is True to avoid exposing data in production.
    """
    if not settings.DEBUG:
        raise HTTPException(status_code=404, detail="Not found")

    resume = db.query(Resume).filter(Resume.id == resume_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    def _proj_link(p):
        return getattr(p, 'github_url', None) or getattr(p, 'demo_url', None) or getattr(p, 'url', None)

    return {
        "id": resume.id,
        "full_name": resume.full_name,
        "email": resume.email,
        "summary": resume.summary,
        "work_experience_count": len(resume.work_experiences) if resume.work_experiences else 0,
        "education_count": len(resume.education) if resume.education else 0,
        "skills": [getattr(s, 'name', str(s)) for s in (resume.skills or [])],
        "projects": [
            {"title": getattr(p, 'title', ''), "link": _proj_link(p)} for p in (resume.projects or [])
        ]
    }


@router.get("/{resume_id}/export_style_debug")
def export_style_debug(resume_id: int, db: Session = Depends(get_db)):
    """Debug endpoint returning the effective template/style values used for export.
    Only available when `settings.DEBUG` is True.
    """
    if not settings.DEBUG:
        raise HTTPException(status_code=404, detail="Not found")

    resume = db.query(Resume).filter(Resume.id == resume_id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    # Try to load template config
    try:
        template = db.query(ResumeTemplate).filter(
            or_(ResumeTemplate.name == (resume.template_id or ''), ResumeTemplate.id == resume.template_id)
        ).first()
    except Exception:
        template = None

    effective = {
        'template_id': resume.template_id,
        'template_name': getattr(template, 'name', None) if template else None,
        'font_family': resume.font_family,
        'accent_color': resume.accent_color,
        'text_color': resume.text_color,
        'heading_color': resume.heading_color,
        'layout': resume.layout,
        'line_spacing': resume.line_spacing,
        'font_size': resume.font_size,
        'heading_size': resume.heading_size,
        'template_config': template.config if template and hasattr(template, 'config') else None
    }

    return { 'id': resume.id, 'effective': effective }


@router.post("/{resume_id}/export")
@router.get("/{resume_id}/export")
def export_resume(resume_id: int, format: str = "pdf", db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Export resume to PDF or DOCX with actual resume data properly formatted."""
    resume = db.query(Resume).filter(Resume.id == resume_id, Resume.user_id == current_user.id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    # Load template config (if any) and compute effective styling values
    try:
        template = db.query(ResumeTemplate).filter(
            or_(ResumeTemplate.name == (resume.template_id or ''), ResumeTemplate.id == resume.template_id)
        ).first()
    except Exception:
        template = None

    effective_font = resume.font_family or 'Roboto'
    effective_accent = resume.accent_color or '#2563eb'
    effective_text = resume.text_color or '#000000'
    effective_heading = resume.heading_color or '#1f2937'
    effective_layout = resume.layout or 'single-column'
    effective_line_spacing = resume.line_spacing or 1.2
    effective_font_size = resume.font_size or 11
    effective_heading_size = resume.heading_size or 14

    if template and template.config and isinstance(template.config, dict):
        cfg = template.config
        effective_font = cfg.get('font_family', effective_font)
        effective_accent = cfg.get('accent_color', effective_accent)
        effective_text = cfg.get('text_color', effective_text)
        effective_heading = cfg.get('heading_color', effective_heading)
        effective_layout = cfg.get('layout', effective_layout)
        effective_line_spacing = cfg.get('line_spacing', effective_line_spacing)
        effective_font_size = cfg.get('font_size', effective_font_size)
        effective_heading_size = cfg.get('heading_size', effective_heading_size)

    fmt = format.lower()
    if fmt not in ("pdf", "docx"):
        raise HTTPException(status_code=400, detail="Unsupported format. Use 'pdf' or 'docx'.")

    tmp_dir = tempfile.gettempdir()
    ts = datetime.utcnow().strftime("%Y%m%d%H%M%S")
    filename = f"resume-{resume.full_name or 'export'}-{ts}.{fmt}"
    filepath = os.path.join(tmp_dir, filename)

    try:
        if fmt == "pdf":
            doc = SimpleDocTemplate(filepath, pagesize=letter, topMargin=0.5*inch, bottomMargin=0.5*inch, leftMargin=0.75*inch, rightMargin=0.75*inch)
            story = []
            styles = getSampleStyleSheet()
            
            title_style = ParagraphStyle(
                'CustomTitle', parent=styles['Heading1'], fontSize=max(14, effective_heading_size + 2), textColor=colors.HexColor(effective_heading), spaceAfter=6
            )
            heading_style = ParagraphStyle(
                'SectionHeading', parent=styles['Heading2'], fontSize=max(10, effective_heading_size), textColor=colors.HexColor(effective_accent), spaceAfter=6, spaceBefore=12
            )
            
            # Header with full name
            if resume.full_name:
                story.append(Paragraph(resume.full_name, title_style))
            
            # Contact info
            contact_info = []
            if resume.email:
                contact_info.append(resume.email)
            if resume.phone:
                contact_info.append(resume.phone)
            if resume.location:
                contact_info.append(resume.location)
            if contact_info:
                story.append(Paragraph(" | ".join(contact_info), styles['Normal']))
            
            # Links
            if resume.linkedin_url or resume.github_url or resume.portfolio_url:
                links = []
                if resume.linkedin_url:
                    links.append(f"LinkedIn: {resume.linkedin_url}")
                if resume.github_url:
                    links.append(f"GitHub: {resume.github_url}")
                if resume.portfolio_url:
                    links.append(f"Portfolio: {resume.portfolio_url}")
                story.append(Paragraph(" | ".join(links), styles['Normal']))
            
            story.append(Spacer(1, 0.2*inch))
            
            # Professional Summary
            if resume.summary:
                story.append(Paragraph("Professional Summary", heading_style))
                story.append(Paragraph(str(resume.summary), styles['Normal']))
                story.append(Spacer(1, 0.15*inch))
            
            # Work Experience
            if resume.work_experiences and len(resume.work_experiences) > 0:
                story.append(Paragraph("Work Experience", heading_style))
                for exp in resume.work_experiences:
                    position_line = f"<b>{exp.position}</b> at {exp.company}"
                    if exp.start_date or exp.end_date:
                        position_line += f" ({exp.start_date or ''} - {exp.end_date or 'Present'})"
                    story.append(Paragraph(position_line, styles['Normal']))
                    if exp.description:
                        story.append(Paragraph(str(exp.description), styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            # Education
            if resume.education and len(resume.education) > 0:
                story.append(Paragraph("Education", heading_style))
                for edu in resume.education:
                    edu_line = f"<b>{getattr(edu, 'degree', '')}</b>"
                    # Education field name is `field_of_study` in the model
                    field_name = getattr(edu, 'field_of_study', None) or getattr(edu, 'field', None)
                    if field_name:
                        edu_line += f" in {field_name}"
                    edu_line += f" from {getattr(edu, 'institution', '')}"
                    story.append(Paragraph(edu_line, styles['Normal']))
                    if hasattr(edu, 'graduation_date') and edu.graduation_date:
                        story.append(Paragraph(f"Graduation: {edu.graduation_date}", styles['Normal']))
                    story.append(Spacer(1, 0.1*inch))
            
            # Skills
            if resume.skills and len(resume.skills) > 0:
                story.append(Paragraph("Skills", heading_style))
                # ResumeSkill model uses `name` for the skill label
                skills_text = ", ".join([getattr(s, 'name', str(s)) for s in resume.skills])
                story.append(Paragraph(skills_text, styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
            
            # Projects
            if resume.projects and len(resume.projects) > 0:
                story.append(Paragraph("Projects", heading_style))
                for proj in resume.projects:
                    proj_line = f"<b>{getattr(proj, 'title', '')}</b>"
                    # Projects may expose github_url or demo_url
                    proj_link = getattr(proj, 'github_url', None) or getattr(proj, 'demo_url', None) or getattr(proj, 'url', None)
                    if proj_link:
                        proj_line += f" ({proj_link})"
                    story.append(Paragraph(proj_line, styles['Normal']))
                    if getattr(proj, 'description', None):
                        story.append(Paragraph(str(getattr(proj, 'description')), styles['Normal']))
                    story.append(Spacer(1, 0.08*inch))

            # Languages (if present)
            languages = getattr(resume, 'languages', None) or getattr(resume, 'resume_languages', None) or []
            if languages and len(languages) > 0:
                story.append(Paragraph("LANGUAGES", heading_style))
                langs_text = ", ".join([getattr(l, 'name', str(l)) + (f" ({getattr(l,'proficiency','')})" if getattr(l,'proficiency',None) else '') for l in languages])
                story.append(Paragraph(langs_text, styles['Normal']))
                story.append(Spacer(1, 0.08*inch))

            # Publications
            publications = getattr(resume, 'publications', None) or getattr(resume, 'resume_publications', None) or []
            if publications and len(publications) > 0:
                story.append(Paragraph("PUBLICATIONS", heading_style))
                for pub in publications:
                    pub_line = f"<b>{getattr(pub,'title','')}</b>"
                    if getattr(pub,'authors',None):
                        pub_line += f" — {getattr(pub,'authors')}"
                    story.append(Paragraph(pub_line, styles['Normal']))
                    if getattr(pub,'publisher',None):
                        story.append(Paragraph(f"Publisher: {getattr(pub,'publisher')}", styles['Normal']))
                    if getattr(pub,'url',None):
                        story.append(Paragraph(f"URL: {getattr(pub,'url')}", styles['Normal']))
                    story.append(Spacer(1, 0.08*inch))

            # Patents
            patents = getattr(resume, 'patents', None) or getattr(resume, 'resume_patents', None) or []
            if patents and len(patents) > 0:
                story.append(Paragraph("PATENTS", heading_style))
                for p in patents:
                    p_line = f"<b>{getattr(p,'title','')}</b>"
                    if getattr(p,'date',None):
                        p_line += f" ({getattr(p,'date')})"
                    story.append(Paragraph(p_line, styles['Normal']))
                    if getattr(p,'description',None):
                        story.append(Paragraph(getattr(p,'description'), styles['Normal']))
                    story.append(Spacer(1, 0.08*inch))

            # References
            references = getattr(resume, 'references', None) or getattr(resume, 'resume_references', None) or []
            if references and len(references) > 0:
                story.append(Paragraph("REFERENCES", heading_style))
                for r in references:
                    ref_line = f"{getattr(r,'name','')} — {getattr(r,'title','')} at {getattr(r,'company','')}"
                    story.append(Paragraph(ref_line, styles['Normal']))
                    if getattr(r,'email',None):
                        story.append(Paragraph(f"Email: {getattr(r,'email')}", styles['Normal']))
                    if getattr(r,'phone',None):
                        story.append(Paragraph(f"Phone: {getattr(r,'phone')}", styles['Normal']))
                    story.append(Spacer(1, 0.08*inch))
            
            doc.build(story)

        else:  # docx
            doc = Document()
            
            if resume.full_name:
                heading = doc.add_heading(resume.full_name, level=1)
                # Align heading according to template/layout preference
                if 'center' in (effective_layout or '').lower() or 'beginner' in (effective_layout or '').lower():
                    heading.alignment = 1
                else:
                    heading.alignment = 0
            
            # Contact info
            contact_info = []
            if resume.email:
                contact_info.append(resume.email)
            if resume.phone:
                contact_info.append(resume.phone)
            if resume.location:
                contact_info.append(resume.location)
            if contact_info:
                p = doc.add_paragraph(" | ".join(contact_info))
                p.alignment = 1
            
            # Links
            if resume.linkedin_url or resume.github_url or resume.portfolio_url:
                links = []
                if resume.linkedin_url:
                    links.append(resume.linkedin_url)
                if resume.github_url:
                    links.append(resume.github_url)
                if resume.portfolio_url:
                    links.append(resume.portfolio_url)
                p = doc.add_paragraph(" | ".join(links))
                p.alignment = 1
            
            doc.add_paragraph()
            
            # Professional Summary
            if resume.summary:
                doc.add_heading("Professional Summary", level=2)
                doc.add_paragraph(str(resume.summary))
            
            # Work Experience
            if resume.work_experiences and len(resume.work_experiences) > 0:
                doc.add_heading("Work Experience", level=2)
                for exp in resume.work_experiences:
                    p = doc.add_paragraph(style='List Bullet')
                    p_run = p.add_run(exp.position)
                    p_run.bold = True
                    p.add_run(f" at {exp.company}")
                    if exp.start_date or exp.end_date:
                        p.add_run(f" ({exp.start_date or ''} - {exp.end_date or 'Present'})")
                    if exp.description:
                        p = doc.add_paragraph(str(exp.description), style='List Bullet 2')
            
            # Education
            if resume.education and len(resume.education) > 0:
                doc.add_heading("Education", level=2)
                for edu in resume.education:
                    p = doc.add_paragraph(style='List Bullet')
                    p_run = p.add_run(getattr(edu, 'degree', ''))
                    p_run.bold = True
                    # Education field name is `field_of_study` in the model
                    field_name = getattr(edu, 'field_of_study', None) or getattr(edu, 'field', None)
                    if field_name:
                        p.add_run(f" in {field_name}")
                    p.add_run(f" from {getattr(edu, 'institution', '')}")
                    if hasattr(edu, 'graduation_date') and edu.graduation_date:
                        p = doc.add_paragraph(f"Graduation: {edu.graduation_date}", style='List Bullet 2')
            
            # Skills
            if resume.skills and len(resume.skills) > 0:
                doc.add_heading("Skills", level=2)
                # ResumeSkill model uses `name` for the skill label
                skills_text = ", ".join([getattr(s, 'name', str(s)) for s in resume.skills])
                doc.add_paragraph(skills_text)
            
            # Projects
            if resume.projects and len(resume.projects) > 0:
                doc.add_heading("Projects", level=2)
                for proj in resume.projects:
                    p = doc.add_paragraph(style='List Bullet')
                    p_run = p.add_run(getattr(proj, 'title', ''))
                    p_run.bold = True
                    # Projects may expose github_url or demo_url or url
                    proj_link = getattr(proj, 'github_url', None) or getattr(proj, 'demo_url', None) or getattr(proj, 'url', None)
                    if proj_link:
                        p.add_run(f" ({proj_link})")
                    if getattr(proj, 'description', None):
                        p = doc.add_paragraph(str(getattr(proj, 'description')), style='List Bullet 2')
            
            doc.save(filepath)

        return FileResponse(
            filepath,
            filename=filename,
            media_type="application/pdf" if fmt == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Export failed: {str(e)}")


class SuggestionRequest(BaseModel):
    sections: Optional[List[str]] = None
    max_suggestions: int = 5


@router.post("/{resume_id}/suggestions")
def resume_suggestions(resume_id: int, payload: SuggestionRequest = None, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Generate resume improvement suggestions using configured LLM provider."""
    payload = payload or SuggestionRequest()

    resume = db.query(Resume).filter(Resume.id == resume_id, Resume.user_id == current_user.id).first()
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")

    # Build a compact prompt from the resume fields
    parts = []
    if resume.full_name:
        parts.append(f"Name: {resume.full_name}")
    if resume.summary:
        parts.append(f"Summary: {resume.summary}")

    if resume.work_experiences:
        parts.append("Work Experience:")
        for exp in resume.work_experiences[:5]:
            parts.append(f"- {exp.position} at {exp.company}: {str(exp.description or '')[:300]}")

    if resume.skills:
        skills_str = ", ".join([getattr(s, 'name', str(s)) for s in resume.skills][:50])
        parts.append(f"Skills: {skills_str}")

    focus = "" if not payload.sections else ("Focus areas: " + ", ".join(payload.sections))
    prompt = "\n".join(parts + [focus, "\nProvide concise, actionable bullet-point suggestions (one per line) to improve this resume. Keep each suggestion short and specific."])

    provider = get_provider()
    try:
        result = provider.generate(prompt, max_tokens=400)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"LLM provider error: {e}")

    if isinstance(result, dict):
        text = result.get('text') or result.get('output') or str(result)
    else:
        text = str(result)

    # Parse suggestions from the text
    suggestions = []
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith(('-', '•', '*')):
            line = line.lstrip('-*• ').strip()
        suggestions.append(line)
        if len(suggestions) >= payload.max_suggestions:
            break

    return {"ok": True, "suggestions": suggestions}
