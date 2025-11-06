"""
Resume Export Service - PDF, DOCX, TXT generation
"""
from fastapi import APIRouter, Depends, HTTPException, Response
from sqlalchemy.orm import Session, joinedload
from typing import Optional
import io
from datetime import datetime

from app.core.db import get_db
from app.core.security import get_current_user
from app.models.user import User
from app.modelsx.resume import (
    Resume, WorkExperience, Education, ResumeProject,
    ResumeSkill, ResumeCertificate, Achievement,
    Language, Publication, Patent, VolunteerWork, Reference
)

router = APIRouter(prefix="/resumes", tags=["Resume Export"])


@router.get("/{resume_id}/export")
async def export_resume(
    resume_id: int,
    format: str = "pdf",  # pdf, docx, txt
    user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """
    Export resume in multiple formats:
    - PDF: Professional formatted document
    - DOCX: Microsoft Word compatible
    - TXT: Plain text version
    """
    # Fetch resume with all relationships
    resume = db.query(Resume).options(
        joinedload(Resume.work_experiences),
        joinedload(Resume.education),
        joinedload(Resume.projects),
        joinedload(Resume.skills),
        joinedload(Resume.certificates),
        joinedload(Resume.achievements)
    ).filter(
        Resume.id == resume_id,
        Resume.user_id == user.id
    ).first()
    
    if not resume:
        raise HTTPException(status_code=404, detail="Resume not found")
    
    # Track download
    resume.downloads += 1
    db.commit()
    
    format = format.lower()
    
    if format == "pdf":
        return await export_pdf(resume, db)
    elif format == "docx":
        return await export_docx(resume, db)
    elif format == "txt":
        return await export_txt(resume, db)
    else:
        raise HTTPException(status_code=400, detail="Unsupported format. Use: pdf, docx, or txt")


def _map_font_to_reportlab(font_family: str) -> tuple[str, str]:
    """Map custom font names to ReportLab built-in fonts.
    Case-insensitive and tolerant of unknown fonts (falls back to Helvetica).
    """
    key = (font_family or "").strip().lower()
    FONT_MAP = {
        "roboto": ("Helvetica", "Helvetica-Bold"),
        "open sans": ("Helvetica", "Helvetica-Bold"),
        "lato": ("Helvetica", "Helvetica-Bold"),
        "montserrat": ("Helvetica", "Helvetica-Bold"),
        "poppins": ("Helvetica", "Helvetica-Bold"),
        "inter": ("Helvetica", "Helvetica-Bold"),
        "playfair display": ("Times-Roman", "Times-Bold"),
        "merriweather": ("Times-Roman", "Times-Bold"),
        "georgia": ("Times-Roman", "Times-Bold"),
        "crimson text": ("Times-Roman", "Times-Bold"),
        "times new roman": ("Times-Roman", "Times-Bold"),
    "century gothic": ("Helvetica", "Helvetica-Bold"),
        "courier new": ("Courier", "Courier-Bold"),
        "source code pro": ("Courier", "Courier-Bold"),
        # generic fallbacks
        "serif": ("Times-Roman", "Times-Bold"),
        "sans-serif": ("Helvetica", "Helvetica-Bold"),
        "monospace": ("Courier", "Courier-Bold"),
    }
    return FONT_MAP.get(key, ("Helvetica", "Helvetica-Bold"))


async def export_pdf(resume: Resume, db: Session):
    """Generate PDF using ReportLab with template styling"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch, mm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib.colors import HexColor, black, white
        from reportlab.pdfgen import canvas as reportlab_canvas
        from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
    except ImportError:
        raise HTTPException(status_code=500, detail="PDF generation not available. Install reportlab.")
    
    # Custom canvas to enforce page limit
    class PageLimitCanvas(reportlab_canvas.Canvas):
        def __init__(self, *args, max_pages=10, **kwargs):
            super().__init__(*args, **kwargs)
            self._max_pages = max_pages
            self._page_count = 0
        
        def showPage(self):
            self._page_count += 1
            if self._page_count >= self._max_pages:
                # Stop rendering after max pages
                return
            super().showPage()
    
    # Map font to ReportLab built-in
    base_font, bold_font = _map_font_to_reportlab(resume.font_family or "Roboto")
    
    # Create PDF buffer
    buffer = io.BytesIO()
    
    # Page setup
    page_size = A4 if resume.page_size == "A4" else letter
    max_pages = resume.max_pages or 10  # Default to 10 pages if not set
    
    doc = SimpleDocTemplate(
        buffer,
        pagesize=page_size,
        leftMargin=resume.page_margins.get("left", 20) * mm if resume.page_margins else 20*mm,
        rightMargin=resume.page_margins.get("right", 20) * mm if resume.page_margins else 20*mm,
        topMargin=resume.page_margins.get("top", 20) * mm if resume.page_margins else 20*mm,
        bottomMargin=resume.page_margins.get("bottom", 20) * mm if resume.page_margins else 20*mm,
    )
    
    # Get colors
    accent_color = HexColor(resume.accent_color or "#2563eb")
    text_color = HexColor(resume.text_color or "#000000")
    heading_color = HexColor(resume.heading_color or "#1f2937")
    
    # Define styles
    styles = getSampleStyleSheet()
    
    # Determine alignment from layout: center for beginner/center layouts, left otherwise
    layout_lower = (resume.layout or "").lower()
    is_centered = "beginner" in layout_lower or "center" in layout_lower
    default_alignment = TA_CENTER if is_centered else TA_LEFT
    
    # Allow override via custom_sections config
    if resume.custom_sections and isinstance(resume.custom_sections, dict):
        align_setting = resume.custom_sections.get('alignment', None)
        if align_setting == 'center':
            default_alignment = TA_CENTER
        elif align_setting == 'left':
            default_alignment = TA_LEFT
        elif align_setting == 'right':
            default_alignment = TA_RIGHT
    
    # Use dynamic heading sizes
    base_font_size = resume.font_size or 11
    heading_font_size = resume.heading_size or 14
    title_font_size = heading_font_size + 4  # Title slightly larger than headings
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=title_font_size,
        textColor=heading_color,
        spaceAfter=6,
        alignment=default_alignment,
        fontName=bold_font
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=heading_font_size,
        textColor=accent_color,
        spaceAfter=6,
        spaceBefore=12,
        fontName=bold_font,
        alignment=default_alignment,
        borderWidth=2 if is_centered else 0,
        borderColor=accent_color if is_centered else black,
        borderPadding=4 if is_centered else 2
    )
    
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=base_font_size,
        textColor=text_color,
        fontName=base_font,
        alignment=default_alignment,
        leading=resume.line_spacing * base_font_size if resume.line_spacing else 14
    )
    
    # Build PDF content
    story = []
    
    # Header - Name and Contact
    # Use full_name if available, otherwise use title, fallback to "Resume"
    display_name = resume.full_name or resume.title or "Resume"
    story.append(Paragraph(display_name, title_style))
    
    contact_info = []
    if resume.email:
        contact_info.append(resume.email)
    if resume.phone:
        contact_info.append(resume.phone)
    if resume.location:
        contact_info.append(resume.location)
    
    if contact_info:
        contact_text = " | ".join(contact_info)
        contact_style = ParagraphStyle('Contact', parent=body_style, alignment=default_alignment, fontSize=10)
        story.append(Paragraph(contact_text, contact_style))
    
    # Links
    links = []
    if resume.linkedin_url:
        links.append(f'<link href="{resume.linkedin_url}">LinkedIn</link>')
    if resume.github_url:
        links.append(f'<link href="{resume.github_url}">GitHub</link>')
    if resume.portfolio_url:
        links.append(f'<link href="{resume.portfolio_url}">Portfolio</link>')
    
    if links:
        links_text = " | ".join(links)
        story.append(Paragraph(links_text, contact_style))
    
    story.append(Spacer(1, 12))
    
    # Professional Summary
    if resume.summary:
        story.append(Paragraph("PROFESSIONAL SUMMARY", heading_style))
        story.append(Paragraph(resume.summary, body_style))
        story.append(Spacer(1, 12))
    
    # Work Experience
    work_experiences = sorted(resume.work_experiences, key=lambda x: x.order_index)
    if work_experiences:
        story.append(Paragraph("EXPERIENCE", heading_style))
        
        for exp in work_experiences:
            # Company and Position
            job_title = f"<b>{exp.position}</b> at {exp.company}"
            if exp.location:
                job_title += f" ({exp.location})"
            story.append(Paragraph(job_title, body_style))
            
            # Dates
            date_range = f"{exp.start_date} - {exp.end_date if not exp.is_current else 'Present'}"
            date_style = ParagraphStyle('DateStyle', parent=body_style, fontSize=10, textColor=HexColor("#6b7280"))
            story.append(Paragraph(date_range, date_style))
            
            # Description
            if exp.description:
                story.append(Paragraph(exp.description, body_style))
            
            # Bullet points
            if exp.bullet_points:
                for bullet in exp.bullet_points:
                    story.append(Paragraph(f"• {bullet}", body_style))
            
            story.append(Spacer(1, 8))
    
    # Education
    education_entries = sorted(resume.education, key=lambda x: x.order_index)
    if education_entries:
        story.append(Paragraph("EDUCATION", heading_style))
        
        for edu in education_entries:
            edu_text = f"<b>{edu.degree}</b>"
            if edu.field_of_study:
                edu_text += f" in {edu.field_of_study}"
            story.append(Paragraph(edu_text, body_style))
            
            institution_text = edu.institution
            if edu.location:
                institution_text += f" ({edu.location})"
            story.append(Paragraph(institution_text, body_style))
            
            date_range = f"{edu.start_date} - {edu.end_date}"
            story.append(Paragraph(date_range, date_style))
            
            if edu.gpa:
                story.append(Paragraph(f"GPA: {edu.gpa}", body_style))
            
            story.append(Spacer(1, 8))
    
    # Skills
    skills = sorted(resume.skills, key=lambda x: x.order_index)
    if skills:
        story.append(Paragraph("SKILLS", heading_style))
        
        # Group skills by category
        skills_by_category = {}
        for skill in skills:
            category = skill.category or "General"
            if category not in skills_by_category:
                skills_by_category[category] = []
            skills_by_category[category].append(skill.name)
        
        for category, skill_names in skills_by_category.items():
            skills_text = f"<b>{category}:</b> {', '.join(skill_names)}"
            story.append(Paragraph(skills_text, body_style))
        
        story.append(Spacer(1, 12))
    
    # Projects
    projects = sorted(resume.projects, key=lambda x: x.order_index)
    if projects:
        story.append(Paragraph("PROJECTS", heading_style))
        
        for project in projects:
            project_title = f"<b>{project.title}</b>"
            story.append(Paragraph(project_title, body_style))
            story.append(Paragraph(project.description, body_style))
            
            if project.tech_stack:
                tech_text = f"<i>Technologies: {', '.join(project.tech_stack)}</i>"
                story.append(Paragraph(tech_text, body_style))
            
            story.append(Spacer(1, 8))
    
    # Certificates
    certificates = sorted(resume.certificates, key=lambda x: x.order_index)
    if certificates:
        story.append(Paragraph("CERTIFICATIONS", heading_style))
        
        for cert in certificates:
            cert_text = f"<b>{cert.name}</b> - {cert.issuing_organization}"
            if cert.issue_date:
                cert_text += f" ({cert.issue_date})"
            story.append(Paragraph(cert_text, body_style))
    
    # Build PDF with page limit enforcement
    def create_canvas(*args, **kwargs):
        return PageLimitCanvas(*args, max_pages=max_pages, **kwargs)
    
    doc.build(story, canvasmaker=create_canvas)
    
    buffer.seek(0)
    
    # Generate filename from full_name or title
    name_for_file = resume.full_name or resume.title or "Resume"
    # Sanitize filename
    safe_name = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in name_for_file)
    filename = f"{safe_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.pdf"
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


async def export_docx(resume: Resume, db: Session):
    """Generate DOCX using python-docx"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(status_code=500, detail="DOCX generation not available. Install python-docx.")
    
    doc = Document()
    
    # Set margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
    
    # Title - Name (use full_name if available, otherwise title, fallback to "Resume")
    display_name = resume.full_name or resume.title or "Resume"
    title = doc.add_heading(display_name, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Contact Info
    contact_parts = []
    if resume.email:
        contact_parts.append(resume.email)
    if resume.phone:
        contact_parts.append(resume.phone)
    if resume.location:
        contact_parts.append(resume.location)
    
    if contact_parts:
        contact = doc.add_paragraph(" | ".join(contact_parts))
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Links
    if any([resume.linkedin_url, resume.github_url, resume.portfolio_url]):
        links = doc.add_paragraph()
        links.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if resume.linkedin_url:
            links.add_run(f"LinkedIn: {resume.linkedin_url} | ")
        if resume.github_url:
            links.add_run(f"GitHub: {resume.github_url} | ")
        if resume.portfolio_url:
            links.add_run(f"Portfolio: {resume.portfolio_url}")
    
    doc.add_paragraph()  # Spacer
    
    # Professional Summary
    if resume.summary:
        doc.add_heading("PROFESSIONAL SUMMARY", level=2)
        doc.add_paragraph(resume.summary)
    
    # Work Experience
    work_experiences = sorted(resume.work_experiences, key=lambda x: x.order_index)
    if work_experiences:
        doc.add_heading("EXPERIENCE", level=2)
        
        for exp in work_experiences:
            job_para = doc.add_paragraph()
            job_para.add_run(f"{exp.position} at {exp.company}").bold = True
            if exp.location:
                job_para.add_run(f" ({exp.location})")
            
            date_para = doc.add_paragraph(
                f"{exp.start_date} - {exp.end_date if not exp.is_current else 'Present'}"
            )
            date_para.style = 'Intense Quote'
            
            if exp.description:
                doc.add_paragraph(exp.description)
            
            if exp.bullet_points:
                for bullet in exp.bullet_points:
                    doc.add_paragraph(bullet, style='List Bullet')
    
    # Education
    education_entries = sorted(resume.education, key=lambda x: x.order_index)
    if education_entries:
        doc.add_heading("EDUCATION", level=2)
        
        for edu in education_entries:
            edu_para = doc.add_paragraph()
            degree_text = edu.degree
            if edu.field_of_study:
                degree_text += f" in {edu.field_of_study}"
            edu_para.add_run(degree_text).bold = True
            
            doc.add_paragraph(f"{edu.institution}" + (f" ({edu.location})" if edu.location else ""))
            doc.add_paragraph(f"{edu.start_date} - {edu.end_date}")
            
            if edu.gpa:
                doc.add_paragraph(f"GPA: {edu.gpa}")
    
    # Skills
    skills = sorted(resume.skills, key=lambda x: x.order_index)
    if skills:
        doc.add_heading("SKILLS", level=2)
        
        skills_by_category = {}
        for skill in skills:
            category = skill.category or "General"
            if category not in skills_by_category:
                skills_by_category[category] = []
            skills_by_category[category].append(skill.name)
        
        for category, skill_names in skills_by_category.items():
            para = doc.add_paragraph()
            para.add_run(f"{category}: ").bold = True
            para.add_run(", ".join(skill_names))
    
    # Projects
    projects = sorted(resume.projects, key=lambda x: x.order_index)
    if projects:
        doc.add_heading("PROJECTS", level=2)
        
        for project in projects:
            proj_para = doc.add_paragraph()
            proj_para.add_run(project.title).bold = True
            doc.add_paragraph(project.description)
            
            if project.tech_stack:
                tech_para = doc.add_paragraph()
                tech_para.add_run("Technologies: ").italic = True
                tech_para.add_run(", ".join(project.tech_stack)).italic = True
    
    # Certificates
    certificates = sorted(resume.certificates, key=lambda x: x.order_index)
    if certificates:
        doc.add_heading("CERTIFICATIONS", level=2)
        
        for cert in certificates:
            cert_para = doc.add_paragraph()
            cert_para.add_run(cert.name).bold = True
            cert_para.add_run(f" - {cert.issuing_organization}")
            if cert.issue_date:
                cert_para.add_run(f" ({cert.issue_date})")
    
    # Save to buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    # Generate filename from full_name or title
    name_for_file = resume.full_name or resume.title or "Resume"
    # Sanitize filename
    safe_name = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in name_for_file)
    filename = f"{safe_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.docx"
    
    return Response(
        content=buffer.getvalue(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )


async def export_txt(resume: Resume, db: Session):
    """Generate plain text version"""
    lines = []
    
    # Header (use full_name if available, otherwise title, fallback to "RESUME")
    display_name = (resume.full_name or resume.title or "RESUME").upper()
    lines.append("=" * 80)
    lines.append(display_name.center(80))
    lines.append("=" * 80)
    lines.append("")
    
    # Contact Info
    contact = []
    if resume.email:
        contact.append(f"Email: {resume.email}")
    if resume.phone:
        contact.append(f"Phone: {resume.phone}")
    if resume.location:
        contact.append(f"Location: {resume.location}")
    
    if contact:
        lines.append(" | ".join(contact).center(80))
        lines.append("")
    
    # Links
    if resume.linkedin_url:
        lines.append(f"LinkedIn: {resume.linkedin_url}")
    if resume.github_url:
        lines.append(f"GitHub: {resume.github_url}")
    if resume.portfolio_url:
        lines.append(f"Portfolio: {resume.portfolio_url}")
    
    if any([resume.linkedin_url, resume.github_url, resume.portfolio_url]):
        lines.append("")
    
    lines.append("-" * 80)
    
    # Professional Summary
    if resume.summary:
        lines.append("")
        lines.append("PROFESSIONAL SUMMARY")
        lines.append("-" * 80)
        lines.append(resume.summary)
        lines.append("")
    
    # Work Experience
    work_experiences = sorted(resume.work_experiences, key=lambda x: x.order_index)
    if work_experiences:
        lines.append("")
        lines.append("EXPERIENCE")
        lines.append("-" * 80)
        
        for exp in work_experiences:
            lines.append("")
            lines.append(f"{exp.position} at {exp.company}")
            if exp.location:
                lines.append(f"Location: {exp.location}")
            lines.append(f"{exp.start_date} - {exp.end_date if not exp.is_current else 'Present'}")
            
            if exp.description:
                lines.append("")
                lines.append(exp.description)
            
            if exp.bullet_points:
                for bullet in exp.bullet_points:
                    lines.append(f"  • {bullet}")
    
    # Education
    education_entries = sorted(resume.education, key=lambda x: x.order_index)
    if education_entries:
        lines.append("")
        lines.append("EDUCATION")
        lines.append("-" * 80)
        
        for edu in education_entries:
            lines.append("")
            degree_text = edu.degree
            if edu.field_of_study:
                degree_text += f" in {edu.field_of_study}"
            lines.append(degree_text)
            lines.append(f"{edu.institution}" + (f" ({edu.location})" if edu.location else ""))
            lines.append(f"{edu.start_date} - {edu.end_date}")
            
            if edu.gpa:
                lines.append(f"GPA: {edu.gpa}")
    
    # Skills
    skills = sorted(resume.skills, key=lambda x: x.order_index)
    if skills:
        lines.append("")
        lines.append("SKILLS")
        lines.append("-" * 80)
        
        skills_by_category = {}
        for skill in skills:
            category = skill.category or "General"
            if category not in skills_by_category:
                skills_by_category[category] = []
            skills_by_category[category].append(skill.name)
        
        for category, skill_names in skills_by_category.items():
            lines.append(f"{category}: {', '.join(skill_names)}")
    
    # Projects
    projects = sorted(resume.projects, key=lambda x: x.order_index)
    if projects:
        lines.append("")
        lines.append("PROJECTS")
        lines.append("-" * 80)
        
        for project in projects:
            lines.append("")
            lines.append(project.title)
            lines.append(project.description)
            
            if project.tech_stack:
                lines.append(f"Technologies: {', '.join(project.tech_stack)}")
    
    # Certificates
    certificates = sorted(resume.certificates, key=lambda x: x.order_index)
    if certificates:
        lines.append("")
        lines.append("CERTIFICATIONS")
        lines.append("-" * 80)
        
        for cert in certificates:
            cert_text = f"{cert.name} - {cert.issuing_organization}"
            if cert.issue_date:
                cert_text += f" ({cert.issue_date})"
            lines.append(cert_text)
    
    lines.append("")
    lines.append("=" * 80)
    lines.append(f"Generated on {datetime.now().strftime('%B %d, %Y')}")
    lines.append("=" * 80)
    
    content = "\n".join(lines)
    
    # Generate filename from full_name or title
    name_for_file = resume.full_name or resume.title or "Resume"
    # Sanitize filename
    safe_name = "".join(c if c.isalnum() or c in (' ', '-', '_') else '_' for c in name_for_file)
    filename = f"{safe_name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}.txt"
    
    return Response(
        content=content.encode('utf-8'),
        media_type="text/plain",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'}
    )
